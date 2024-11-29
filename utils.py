# Utility Functions for File Conversion

from io import BytesIO
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF for PDF handling

# Function to convert response to TXT
def convert_to_txt(response):
    return str(response).encode('utf-8', errors='ignore')

# Function to convert response to DOCX
def convert_to_docx(response):
    doc = Document()
    doc.add_heading("Prompt", 1)
    doc.add_paragraph(response)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream